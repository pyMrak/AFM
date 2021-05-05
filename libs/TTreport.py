# -*- coding: utf-8 -*-
"""
Created on Mon Apr 29 10:31:06 2019

@author: andmra2
"""

from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import openpyxl
from matplotlib import pyplot as plt
from numpy import array, where, transpose, linspace
# from pandas.compat import StringIO
from io import BytesIO

from os import system, getcwd, listdir, startfile
from os.path import isdir, isfile
from scipy.stats import norm
from PIL import Image, ImageFont, ImageDraw
from datetime import datetime
from time import time

from libs.data_manipulation import rearrange_up
from libs.logger import messageLogger


# from subprocess import check_call


class TTreportBuilder(object):
    months = {'01': 'JAN', '02': 'FEB', '03': 'MAR', '04': 'APR', '05': 'MAJ', '06': 'JUN',
              '07': 'JUL', '08': 'AVG', '09': 'SEP', '10': 'OKT', '11': 'NOV', '12': 'DEC'}

    def __init__(self, settings, gdat=None):
        self.figN = 0
        self.tablesData = []
        self.settings = settings
        self.picFont = ImageFont.truetype("calibri.ttf", 80)
        self.lastFile = None
        self.gdat = gdat
        self.messageLogger = messageLogger()

    def defineStyles(self, document):
        # document.add_heading('Document Title', 0)
        document.styles.add_style('Naslov', WD_STYLE_TYPE.PARAGRAPH)
        document.styles['Naslov'].font.name = 'Calibri'
        document.styles['Naslov'].font.size = Pt(16)
        document.styles['Naslov'].font.bold = True

        document.styles.add_style('Podnaslov', WD_STYLE_TYPE.PARAGRAPH)
        document.styles['Podnaslov'].font.name = 'Calibri'
        document.styles['Podnaslov'].font.size = Pt(11)
        document.styles['Podnaslov'].font.bold = True

    def addIzvedbe(self, document):
        izvedbaSvecke = document.add_paragraph('Izvedba svečke: ', style='Podnaslov')
        izvedbaSvecke.add_run(self.settings.izvedbaSkupno).bold = False
        if len(self.settings.groups) > 1:
            for group, sn in zip(self.settings.groups, self.sampleNr):
                if self.settings.enakaIzvedba:
                    izvedbaSvecke.add_run('\n\t\t' + group + ': ' + str(sn) + 'kos').bold = False
                    # print('1:',self.settings.groups, self.settings.enakaIzvedba)
                else:
                    # print('2:',self.settings.groups, self.settings.enakaIzvedba)
                    izvedbaSvecke.add_run(
                        '\n\t\t' + group + ': ' + self.settings.izvedba[group] + ' - ' + str(sn) + 'kos').bold = False
        elif len(self.settings.groups) == 1:
            # print('3:',self.settings.groups)
            izvedbaSvecke.add_run(': ' + str(self.sampleNr[0]) + 'kos').bold = False

    def addProcedura(self, document):
        proceduraTT = document.add_paragraph('Procedura TT: ', style='Podnaslov')
        procedure = self.settings.proceduraTT
        # proceduraTT.add_run(procedura[0]).bold = False
        if self.settings.enakaProcedura:
            procedura = self.settings.proceduraTT[self.settings.groups[0]].split('\n')
            for pro in procedura:
                proceduraTT.add_run(pro + '\n').bold = False
        else:
            for i, group in enumerate(self.settings.groups):
                if i:
                    t = '\t'
                else:
                    t = ''
                procedura = procedure[group].split('\n')
                proceduraTT.add_run(t + '\t' + group + ': ' + procedura[0] + '\n').bold = False
                for pro in procedura[1:]:
                    proceduraTT.add_run('\t\t ' + pro + '\n').bold = False

    def isOK(self, val):
        if val == 'OK' or val == 'Ok' or val == 'ok' or val == 'oK':
            return True
        else:
            False

    def getTableData(self, filename):
        book = openpyxl.load_workbook(filename, data_only=True)
        sheet = book['Tabela']
        tableData = []
        nCol = 30
        for r in range(300):
            if sheet.cell(row=r + 1, column=4).value:
                if self.isOK(sheet.cell(row=r + 1, column=1).value) or self.settings.vkljuciVseGP or r == 0:
                    tableData.append([])
                    for c in range(nCol):
                        data = sheet.cell(row=r + 1, column=c + 4).value
                        if r:
                            tableData[-1].append(data)
                        else:
                            if data:
                                tableData[-1].append(data)
                            else:
                                nCol = c
                                break
            else:
                break
        book.close()
        return array(tableData)

    def getTablesData(self, paths):
        fileNames = []
        self.tablesData = []
        self.sampleNr = []
        for path in paths:
            ld = listdir(path)
            appended = False
            for doc in ld:
                if doc[-8:] == 'AFM.xlsx':
                    fileNames.append(path + '/' + doc)
                    appended = True
            if not appended:
                groupName = path.split('/')[-1]
                self.messageLogger.writeWarning(
                    'Skupina ' + groupName + ' nima datoteke s fun. meritvami oz je njeno ime nepravilno.', self.gdat,
                    False)
                fileNames.append(None)
        for fileName in fileNames:
            if fileName:
                self.tablesData.append(self.getTableData(fileName))
                self.sampleNr.append(len(self.tablesData[-1]) - 1)
            else:
                self.tablesData.append(None)
                self.sampleNr.append(0)

    def createFuncTables(self, document):
        for tableData, group in zip(self.tablesData, self.settings.groups):
            if tableData is not None:
                self.createTable(document, tableData[:, :6],
                                 title='Tabela s funkcionalnimi karakeristikami skupine ' + group + '.')

    def createFailureTables(self, document):
        tables = 0
        for tableData, group in zip(self.tablesData, self.settings.groups):
            if tableData is not None:
                if 'odpoved' in tableData[0, -1].lower():
                    if tableData is not None:
                        if len(self.settings.groups) > 1:
                            self.createTable(document, array([tableData[:, 0], tableData[:, -1]]),
                                             title='Tabela s cikli odpovedi skupine ' + group + '.', integer=True)
                            tables += 1
                        else:
                            self.createTable(document, array([tableData[:, 0], tableData[:, -1]]),
                                             title='Tabela s cikli odpovedi.', integer=True)
                            tables += 1
                else:
                    self.messageLogger.writeWarning('Skupina ' + group + ' nima definiranih odpovedi TT', self.gdat,
                                                    False)
        return tables

    def createTable(self, document, tableData, title='', integer=False):
        # color_heading = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        # color_1_1row = parse_xml(r'<w:shd {} w:fill="c7d5e0"/>'.format(nsdecls('w')))
        # color_1_row = parse_xml(r'<w:shd {} w:fill="a0c6e5"/>'.format(nsdecls('w')))
        # color_2_1row = parse_xml(r'<w:shd {} w:fill="9aaab7"/>'.format(nsdecls('w')))
        # color_2_row = parse_xml(r'<w:shd {} w:fill="72a4cc"/>'.format(nsdecls('w')))

        document.add_paragraph(title, style='Podnaslov')
        document.paragraphs[-1].paragraph_format.space_after = 100
        tableStyle = [
            [parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w'))) for i in range(len(tableData[0]))]]
        for i in range(len(tableData) - 1):
            if i % 2:
                tableStyle.append([parse_xml(r'<w:shd {} w:fill="9aaab7"/>'.format(nsdecls('w')))])
            else:
                tableStyle.append([parse_xml(r'<w:shd {} w:fill="c7d5e0"/>'.format(nsdecls('w')))])
            for j in range(len(tableData[0]) - 1):
                if i % 2:
                    tableStyle[-1].append(parse_xml(r'<w:shd {} w:fill="72a4cc"/>'.format(nsdecls('w'))))
                else:
                    tableStyle[-1].append(parse_xml(r'<w:shd {} w:fill="a0c6e5"/>'.format(nsdecls('w'))))

        table = document.add_table(rows=len(tableData), cols=len(tableData[0]))
        table.style = 'TableGrid'
        for r, row in enumerate(tableData):
            for c, data in enumerate(row):
                try:
                    if c == 0:
                        table.rows[r].cells[c].text = str(data)
                    elif c == 2 or c == 3 or integer:
                        table.rows[r].cells[c].text = str(int(round(float(data), 0)))  # .replace('.',',')
                    else:
                        table.rows[r].cells[c].text = str(round(float(data), 2)).replace('.', ',')
                except:
                    if data is not None:
                        table.rows[r].cells[c].text = str(data)
                    else:
                        table.rows[r].cells[c].text = '/'
                table.rows[r].cells[c]._tc.get_or_add_tcPr().append(tableStyle[r][c])
        document.add_paragraph()
        document.paragraphs[-1].paragraph_format.space_after = 100

    def addBoxPlots(self, document):
        if len(self.settings.groups) == 1:
            if self.tablesData[0] is not None:
                title = 'Čas dosega'
                groups = [self.tablesData[0][0, 1]]
                self.addBoxPlot(document, array(self.tablesData[0][1:, 1], dtype='float64'), groups, title, 284)
                title = 'Temperaturne karakteristike'
                groups = self.tablesData[0][0, 2:4]
                self.addBoxPlot(document, array(self.tablesData[0][1:, 2:4], dtype='float64'), groups, title, 284)
                title = 'Tokovne karakteristike'
                groupsX = self.tablesData[0][0, 4:5]
                groupsY = self.tablesData[0][0, 5:6]
                self.addTwoBoxPlots(document, array(self.tablesData[0][1:, 4:5], dtype='float64'),
                                    array(self.tablesData[0][1:, 5:6], dtype='float64'), groupsX, groupsY, title, 284)
        else:
            itr = None
            for tableData in self.tablesData:
                if tableData is not None:
                    itr = tableData[0, 1:-1]
                    break
            if itr is not None:
                for i, title in enumerate(itr):
                    lis = []
                    for tableData in self.tablesData:
                        if tableData is not None:
                            lis.append(array(tableData[1:, i + 1], dtype='float64'))
                        else:
                            lis.append(array([]))
                    self.addBoxPlot(document, lis, self.settings.groups, title, 284)

    def addBoxPlot(self, document, x, labels, title, width):
        # x = [[1,5,7,3,8,5], [9, 6,5,7,4,5,3,5]]
        plt.figure(self.figN, figsize=(6, 2.85), frameon=False)
        self.figN += 1
        plt.boxplot(x, labels=labels, patch_artist=True, medianprops={
            'color': 'orange'})  # , notch=None, sym=None, vert=None, whis=None, positions=None, widths=None, patch_artist=None, bootstrap=None, usermedians=None, conf_intervals=None, meanline=None, showmeans=None, showcaps=None, showbox=None, showfliers=None, boxprops=None, labels=None, flierprops=None, medianprops=None, meanprops=None, capprops=None, whiskerprops=None, manage_ticks=True, autorange=False, zorder=None, *, data=None)[source]
        plt.title(title)
        plt.grid(color='#878787')
        memfile = BytesIO()  # StringIO()
        plt.savefig(memfile, format='png')
        plt.close()
        document.add_picture(memfile)  # , width=Pt(width))

    def addTwoBoxPlots(self, document, x, y, labelsX, labelsY, title, width):
        # x = [[1,5,7,3,8,5], [9, 6,5,7,4,5,3,5]]
        # plt.figure(self.figN, figsize=(6,2.85))
        # plt.figure(self.figN, figsize=(6,2.85))
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(6, 2.85), frameon=False)
        fig.number = self.figN
        # fig.set_figheight(2.85)
        # fig.set_figwidth(6)
        # plt.figure(self.figN)
        self.figN += 1
        ax1.boxplot(x, labels=labelsX, patch_artist=True, medianprops={
            'color': 'orange'})  # , notch=None, sym=None, vert=None, whis=None, positions=None, widths=None, patch_artist=None, bootstrap=None, usermedians=None, conf_intervals=None, meanline=None, showmeans=None, showcaps=None, showbox=None, showfliers=None, boxprops=None, labels=None, flierprops=None, medianprops=None, meanprops=None, capprops=None, whiskerprops=None, manage_ticks=True, autorange=False, zorder=None, *, data=None)[source]
        ax2.boxplot(y, labels=labelsY, patch_artist=True, medianprops={'color': 'orange'})
        fig.suptitle(title)
        ax1.grid(color='#878787')
        ax2.grid(color='#878787')
        memfile = BytesIO()  # StringIO()
        plt.savefig(memfile, format='png')
        plt.close()
        document.add_picture(memfile)  # , width=Pt(width))
        # document.add_paragraph()

    def createFailureGauss(self, document):
        plt.style.use('ggplot')
        histColors = ['#4ba3eb', '#84e09d', '#f5c076', '#e0dd7b', '#93d9d1', '#bc7ec2', '#f06c6c', '#cfeb88']
        gaussColors = ['#02589e', '#308046', '#db8000', '#b8b200', '#00ab97', '#6f1078', '#b30404', '#6d940c']
        fontColors = ['#003c6e', '#266637', '#ab6400', '#8f8a00', '#00786a', '#500a57', '#7d0000', '#547309']
        plt.figure(self.figN, figsize=(6, 3), frameon=False)

        # Generate a normal distribution, center at x=0 and y=5
        stats = []
        maxC = 0
        minC = 1e10
        i = 0
        j = 0
        data = []
        colors = []
        edgeColors = []
        binSpan = []
        groups = []
        for k, tableData in enumerate(self.tablesData):
            if tableData is not None:
                groups.append(self.settings.groups[k])
                j += 1
                notNoneIdx = where(tableData[1:, -1] != None)[0]
                if len(notNoneIdx) > 0:
                    data.append(array(tableData[1:, -1][notNoneIdx], dtype='float64'))
                    colors.append(histColors[i])
                    edgeColors.append(gaussColors[i])
                    if len(data[-1]) < 8:
                        n_bins = 5
                    else:
                        n_bins = 7
                    # if len(data[-1]) > 0:
                    sigma = data[-1].std()
                    mean = data[-1].mean()
                    binSpan.append((max(data[-1]) - min(data[-1])) / n_bins)
                    # else:
                    #     sigma = None
                    #     mean = None
                    #     binSpan.append(500)
                    stats.append([mean, sigma])
                    # plt.hist(data, bins=n_bins, label=group, density=True, color=histColors[i], alpha=0.75,  edgecolor=gaussColors[i], linewidth=1.2)

                    if i == 0:
                        maxC = max(data[-1])
                        minC = min(data[-1])
                    else:
                        if maxC < max(data[-1]):
                            maxC = max(data[-1])
                        if minC > min(data[-1]):
                            minC = min(data[-1])
                    i += 1
                    if i >= len(histColors):
                        i = 0
                else:
                    data.append(None)
                    colors.append(None)
                    edgeColors.append(None)
                    stats.append([None, None])
                    binSpan.append(0)
        spanC = maxC - minC
        minX = minC - spanC * 0.1
        maxX = maxC + spanC * 0.1

        # print(minC, maxC, minX, maxX)

        if len(binSpan) > 0:
            n_bins = int(round(spanC / (sum(binSpan) / len(binSpan)), 0))
            binWidth = spanC / n_bins
        else:
            n_bins = 0
            binWidth = 1
        # print(binSpan, spanC, sum(binSpan)/len(binSpan) , n_bins)
        x = linspace(minX, maxX, 1000)

        for i, d in enumerate(data):
            if d is not None:
                plt.hist(d, bins=n_bins, label=groups[i], density=True, color=colors[i],
                         alpha=0.5, range=(minC, maxC), edgecolor=edgeColors[i], linewidth=1.2)

        i = 0
        for stat in stats:
            if stat[1] is not None:
                gauss = norm.pdf(x, stat[0], stat[1])
                txtYPos = max(gauss) * 1.02
                offset = max(gauss) * 0.15
                plt.plot(x, gauss, color=gaussColors[i])
                i += 1
                if i >= len(gaussColors):
                    i = 0
        i = 0
        yMaxPlot = 0
        usedPositions = []
        graphSpan = maxX - minX
        for stat in stats:
            if stat[0] is not None:
                meanTxt = 'mean: ' + str(int(round(stat[0], 0)))
                stdTxt = 'std: ' + str(round(stat[1], 1))
                yMin, yMax = plt.ylim()
                ran = yMax - yMin
                if stat[1] > 0:
                    gauss = norm.pdf(x, stat[0], stat[1])
                    txtYPos = max(gauss) + ran * 0.02
                else:
                    txtYPos = 1 / binWidth + ran * 0.02
                offset = ran * 0.07
                clearOffset = 0
                notFinished = True
                while notFinished:
                    notFinished = False
                    for usedPosition in usedPositions:
                        if (usedPosition[0][0] < stat[0] < usedPosition[0][1] and
                                usedPosition[1][0] < txtYPos + clearOffset < usedPosition[1][1]):
                            clearOffset += usedPosition[1][1] - (txtYPos + clearOffset)
                            notFinished = True
                if txtYPos + 2 * offset + clearOffset > yMaxPlot:
                    # plt.ylim(yMin, txtYPos + 2 * offset)
                    yMaxPlot = txtYPos + 2 * offset + clearOffset

                plt.text(stat[0], txtYPos + offset + clearOffset, meanTxt, color=fontColors[i])
                plt.text(stat[0], txtYPos + clearOffset, stdTxt, color=fontColors[i])
                usedPositions.append([[stat[0] - graphSpan * 0.22, stat[0] + graphSpan * 0.22],
                                      [txtYPos + clearOffset - 2 * offset, txtYPos + 2 * offset + clearOffset]])
                i += 1
                if i >= len(fontColors):
                    i = 0
        print('used:', usedPositions)
        plt.xlim(minX, maxX)
        plt.ylim(0, yMaxPlot)
        plt.legend()
        plt.grid(color='#878787')
        plt.title('Histogram odpovedi')
        memfile = BytesIO()  # StringIO()
        plt.savefig(memfile, format='png')
        plt.close()
        self.figN += 1
        document.add_picture(memfile)

    def addRTG(self, document):
        for i, group in enumerate(self.settings.groups):
            if self.tablesData[i] is not None:
                if len(self.settings.groups) > 1:
                    self.addRTGpictures(document, group)
                else:
                    self.addRTGpictures(document)

    def addRTGpictures(self, document, group=''):
        nrPerRow = self.settings.RTGperRow
        title = 'RTG analiza odpovedi ' + group
        document.add_paragraph(title, style='Podnaslov')

        # document.paragraphs[-1].paragraph_format.space_after = 100
        texts = []
        if group == '':
            i = 0
            sampleNames = self.tablesData[0][1:, 0]
        else:
            i = self.settings.groups.index(group)
            sampleNames = self.tablesData[i][1:, 0]
        pasted = False
        for j, sampleName in enumerate(sampleNames):
            if isfile(self.settings.folders[i] + '/RTG/' + sampleName + '.jpg'):
                file = self.settings.folders[i] + '/RTG/' + sampleName + '.jpg'
                texts.append(sampleName)
            elif isfile(self.settings.folders[i] + '/RTG/' + sampleName.split('_')[0] + '.jpg'):
                file = self.settings.folders[i] + '/RTG/' + sampleName.split('_')[0] + '.jpg'
                texts.append(sampleName)
            else:
                file = None
                texts.append('')
                # print('WARNING!: )
                self.messageLogger.writeWarning('RTG slika za vzorec ' + sampleName + ' ne obstaja.', self.gdat, False)
            if j % nrPerRow == 0:
                RTGimage = Image.new('RGB', (nrPerRow * 1000 + 200 * (nrPerRow - 1), 1000), color=(255, 255, 255, 0))
                offset = 0
            if file:
                RTGimage.paste(Image.open(file), (offset, 0))
                pasted = True

                # draw.text((x, y),"Sample Text",(r,g,b))

            offset += 1200
            if j % nrPerRow == nrPerRow - 1:
                self.addRTGPicture(document, RTGimage, texts)
                texts = []
                pasted = False
        if pasted:
            self.addRTGPicture(document, RTGimage, texts)

    def addRTGPicture(self, document, RTGimage, texts):
        drawText = ImageDraw.Draw(RTGimage)
        for k, t in enumerate(texts):
            # print('texts: ', t)
            drawText.text((k * 1200 + 500, 10), t, (0, 0, 0), font=self.picFont)  # k*1200+500
        memfile = BytesIO()
        RTGimage.save(memfile, format='png')
        document.add_picture(memfile, width=Inches(6))

    def getLastFile(self):
        if self.lastFile:
            return [self.lastFile]
        else:
            return []

    def createReport(self, openFile=False):
        # preberi podatke iz excela
        start = time()
        self.messageLogger.writeProgressInfo('berem podatke o TT...', self.gdat, False)
        self.getTablesData(self.settings.folders)
        if len(self.tablesData) == 0 or all(td is None for td in self.tablesData):
            return []
        else:
            # opri nov wordov dokument
            document = Document()
            self.defineStyles(document)
            self.messageLogger.writeProgressInfo('Ustvarjam glavo dokumenta...', self.gdat, False)
            # dodaj naslove
            naslov = document.add_paragraph('POROČILO O IZVEDENIH TRAJNOSTNIH TEKIH SVEČK ŽARILNIH', style='Naslov')
            imeTT = document.add_paragraph('TT-' + str(self.settings.TTnr), style='Podnaslov')

            # dodaj podatke o TT
            namen = document.add_paragraph('Namen testa: ', style='Podnaslov')
            namen.add_run(self.settings.namenTesta).bold = False
            kodaSvecke = document.add_paragraph('Svečka: ', style='Podnaslov')
            kodaSvecke.add_run(self.settings.kodaSvecke).bold = False
            self.addIzvedbe(document)
            self.addProcedura(document)
            nrCiklov = document.add_paragraph('Število ciklov: ', style='Podnaslov')
            nrCiklov.add_run(self.settings.cikli).bold = False

            # dodaj poglavje s funkcionalnimi meritvami
            self.messageLogger.writeProgressInfo('Ustvarjam poglavje s funkcionalnimi meritvami...', self.gdat, False)
            funkcMer = document.add_heading('1. Funkcionalne meritve', 1)
            document.add_paragraph(' ')
            self.createFuncTables(document)
            self.addBoxPlots(document)

            # dodaj poglavje z analizo odpovedi
            self.messageLogger.writeProgressInfo('Ustvarjam poglavje z odpovedmi...', self.gdat, False)
            odpovedi = document.add_heading('2. Odpovedi', 1)

            tablesNr = self.createFailureTables(document)
            if tablesNr > 0:
                self.createFailureGauss(document)
                self.addRTG(document)
                document.add_paragraph('\nMehanizem odpovedi\n', style='Podnaslov')
            else:
                document.add_paragraph('\nNi odpovedi.\n', style='Podnaslov')

            # dodaj naslov za zaključek
            self.messageLogger.writeProgressInfo('Zaključujem dokument...', self.gdat, False)
            zaključek = document.add_heading('3. Komentar rezultatov', 1)

            # dodaj avtorja in datum
            now = datetime.now()  # current date and time
            date = now.strftime("%d.%m.%Y")
            month = self.months[now.strftime("%m")]
            year = date[-2:]
            document.add_paragraph('\n\n')
            pripravil = document.add_paragraph('Pripravil/a:\n', style='Podnaslov')
            pripravil.add_run(
                self.settings.author + ', ' + date).bold = False  # document.add_paragraph(avtor, style='Podnaslov').bold = False

            # dodaj meta podatke dokumenta
            core_properties = document.core_properties
            core_properties.author = 'AFM by Andrej Mrak \xa9'
            core_properties.company = 'Hidria d.o.o.'

            # shrani dokument v izbrano mapo
            self.messageLogger.writeProgressInfo('Shranjujem dokument...', self.gdat, False)
            saved = False
            docFileName = self.settings.path + '/POROČILO O IZVEDENIH TRAJNOSTNIH TEKIH SVEČK ŽARILNIH TT-' + self.settings.TTnr + ' ' + month + year + '.docx'
            for i in range(10):
                try:
                    document.save(docFileName)
                    saved = True
                    break
                except:
                    docFileName = self.settings.path + '/POROČILO O IZVEDENIH TRAJNOSTNIH TEKIH SVEČK ŽARILNIH TT-' + self.settings.TTnr + ' ' + month + year + '(' + str(
                        i + 1) + ').docx'
            if not saved:
                self.messageLogger.writeError('Dokumenta iz neznanega razloga ni mogoče shraniti.', self.gdat, False)
            elif i:
                self.messageLogger.writeWarning(
                    'Dokument z istim imenom je že odprt. Nova verzija dokumenta shranjena s končnico (' + str(
                        i) + ').', self.gdat, False)
            # system("start " +getcwd()+"/demo.docx")
            # print(docFileName, isfile(docFileName))
            # check_call(['open', docFileName])

            # odpri dokument
            if openFile:
                # system('start "'+getcwd()+'/'+docFileName+'"')
                startfile(docFileName.replace('/', '\\'))
            self.lastFile = docFileName
            end = time()
            dur = str(round(end - start, 2)) + 's.'
            # vrni lokacijo shranjenega dokumenta
            self.messageLogger.writeProgressInfo('Dokument shranjen. Obdelava trajala ' + dur, self.gdat, False)
            return [docFileName]

        # print(array(tableData))

        # for i, title in enumerate(tableData[0,1:]):
        #    self.addBoxPlot(document, array(tableData[1:,i+1], dtype='float64'), ['group'], title, 284)


#        document.add_heading('Heading, level 1', level=1)
#        document.add_paragraph('Intense quote', style='Intense Quote')
#        
#        document.add_paragraph(
#            'first item in unordered list', style='List Bullet'
#        )
#        document.add_paragraph(
#            'first item in ordered list', style='List Number'
#        )
#        
#        document.add_picture('monty-truth.png', width=Cm(1.25))
#        
#        document.add_page_break()


class TTreportSettings(object):

    def __init__(self, login=None, gdat=None):
        self.izvedbaSkupno = ''
        self.izvedba = {}
        self.TTnr = ''
        self.namenTesta = ''
        if login:
            self.author = login.username
        else:
            self.author = ''
        self.kodaSvecke = ''
        self.proceduraTT = {}
        self.enakaProcedura = True
        self.enakaIzvedba = True
        self.cikli = ''
        self.vkljuciVseGP = False
        self.RTGperRow = 3
        self.groups = []
        self.folders = []
        self.path = ''
        self.gdat = gdat
        self.messageLogger = messageLogger()

    def setFolder(self, path):
        self.groups = []
        self.folders = []
        self.path = path
        if isdir(path):
            for doc in rearrange_up(listdir(path)):
                # print(doc)
                folder = path + '/' + doc
                if isdir(folder):
                    self.groups.append(doc)
                    self.folders.append(folder)
                # print(doc[-7:], 'AFM.xls', doc[-7:] == 'AFM.xls')
                if doc[-7:] == 'AFM.xls':
                    self.groups = [doc[:-7]]
                    self.folders = [path]
                    break
            # print('setFolder:',self.groups, self.folders)
            self.TTnr = path.split('/')[-1].split('-')[0].strip()
            self.izvedba = {group: '' for group in self.groups}
            self.proceduraTT = {group: '' for group in self.groups}
        return self.groups

    def setProcedura(self, group, value):
        # print('enaka procedura:', self.enakaProcedura)
        if not self.enakaProcedura:
            if group in self.proceduraTT:
                self.proceduraTT[group] = value
        else:
            for group in self.proceduraTT:
                self.proceduraTT[group] = value
        # print(self.proceduraTT)

    def setIzvedbaSkupno(self, izvedba):
        self.izvedbaSkupno = izvedba

    def setIzvedba(self, group, value):
        if not self.enakaIzvedba:
            if group in self.izvedba:
                self.izvedba[group] = value
        else:
            for group in self.izvedba:
                self.izvedba[group] = value

    def setEnakaIzvedba(self, tf):
        self.enakaIzvedba = bool(tf)
        if self.enakaIzvedba:
            for group in self.izvedba:
                if self.izvedba[group] != '':
                    self.setIzvedba(group, self.izvedba[group])
                    break

    def setEnakaProcedura(self, tf):
        self.enakaProcedura = bool(tf)
        if self.enakaProcedura:
            for group in self.proceduraTT:
                if self.proceduraTT[group] != '':
                    self.setProcedura(group, self.proceduraTT[group])
                    break

    def setKodaSvecke(self, koda):
        self.kodaSvecke = koda

    def setAuthor(self, author):
        self.author = author

    def setNamenTesta(self, namen):
        self.namenTesta = namen

    def setCikli(self, cikli):
        try:
            cikli = int(cikli)
            if cikli < 0:
                self.messageLogger.writeWarning('Nastavljeni cikli niso naravno število.', self.gdat,
                                                False)  # print('WARNING!: Nastavljeni cikli niso naravno število.')
        except:
            pass  # print('WARNING!: Nastavljeni cikli niso število.')
        self.cikli = str(cikli)

    def setRTGperRow(self, nr):
        self.RTGperRow = nr

    def getEnakaIzvedba(self):
        return self.enakaIzvedba

    def getEnakaProcedura(self):
        return self.enakaProcedura

    def getProcedura(self, group):
        if group in self.proceduraTT:
            return self.proceduraTT[group]
        else:
            self.messageLogger.writeWarning('Skupina ' + group + ' ne obstaja', self.gdat,
                                            False)  # print('WARNING!: Skupina '+group+' ne obstaja')
            return ''

    def getIzvedba(self, group):
        if group in self.izvedba:
            return self.izvedba[group]
        else:
            self.messageLogger.writeWarning('Skupina ' + group + ' ne obstaja', self.gdat,
                                            False)  # print('WARNING!: Skupina '+group+' ne obstaja')
            return ''

    def getIzvedbaSkupno(self):
        return self.izvedbaSkupno

    def getCikli(self):
        return self.cikli

    def getKodaSvecke(self):
        return self.kodaSvecke

    def getAuthor(self):
        return self.author

    def getNamenTesta(self):
        return self.namenTesta

    def getRTGperRow(self):
        return self.RTGperRow


class TTreport(object):

    def __init__(self, gdat=None, login=None):
        self.settings = TTreportSettings(login, gdat)
        self.reportBuilder = TTreportBuilder(self.settings, gdat)

    def setFolder(self, path):
        return self.settings.setFolder(path)

    def setEnakaIzvedba(self, tf):
        self.settings.setEnakaIzvedba(tf)

    def setEnakaProcedura(self, tf):
        self.settings.setEnakaProcedura(tf)

    def setProcedura(self, group, value):
        self.settings.setProcedura(group, value)

    def setIzvedba(self, group, value):
        self.settings.setIzvedba(group, value)

    def setIzvedbaSkupno(self, izvedba):
        self.settings.setIzvedbaSkupno(izvedba)

    def setCikli(self, cikli):
        self.settings.setCikli(cikli)

    def setKodaSvecke(self, koda):
        self.settings.setKodaSvecke(koda)

    def setAuthor(self, author):
        self.settings.setAuthor(author)

    def setNamenTesta(self, namen):
        self.settings.setNamenTesta(namen)

    def setRTGperRow(self, nr):
        self.settings.setRTGperRow(nr)

    def getEnakaIzvedba(self):
        return self.settings.getEnakaIzvedba()

    def getEnakaProcedura(self):
        return self.settings.getEnakaProcedura()

    def getProcedura(self, group):
        return self.settings.getProcedura(group)

    def getIzvedba(self, group):
        return self.settings.getIzvedba(group)

    def getIzvedbaSkupno(self):
        return self.settings.getIzvedbaSkupno()

    def getCikli(self):
        return self.settings.getCikli()

    def getKodaSvecke(self):
        return self.settings.getKodaSvecke()

    def getAuthor(self):
        return self.settings.getAuthor()

    def getNamenTesta(self):
        return self.settings.getNamenTesta()

    def getRTGperRow(self):
        return self.settings.getRTGperRow()

    def getLastFile(self):
        return self.reportBuilder.getLastFile()

    def createReport(self):
        return self.reportBuilder.createReport()


if __name__ == "__main__":
    izvedbe = ["Cevka INCONEL 602, Lasersko varjenje, 4 skupine uporov:",
               "plazemska obdelava (parametri neznani)",
               "plazemska obdelava (parametri neznani)",
               "plazemska obdelava (parametri neznani)",
               "referenca (neobdelan upor)"]

    procedura = '2s-11V; 18s-6,1V; hlajenje 100s-off\n2s-11V; 18s-6,1V; hlajenje 30s-off'

    TTR = TTreport()
    # 1 group '188 - VW Rekvalifikacija'
    # 3 groups '187 - VW Failure Curves'
    multiple = 0
    if multiple:
        TTR.setFolder('187 - VW Failure Curves')
    else:
        TTR.setFolder('188 - VW Rekvalifikacija')
    TTR.setIzvedbaSkupno(izvedbe[0])
    for group, izvedba in zip(TTR.settings.groups, izvedbe[1:]):
        TTR.setIzvedba(group, izvedba)
    TTR.setProcedura(None, procedura)
    TTR.setCikli('Do odpovedi.')
    TTR.setAuthor('Andrej Mrak')
    TTR.setNamenTesta('Validacija upora.')
    TTR.setKodaSvecke('5011-721-968')

    # print(TTR.settings.groups)
    # print(TTR.settings.folders)
    # print(TTR.settings.TTnr)
    TTR.createReport(True)
